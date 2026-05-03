import mimetypes
import os
from pathlib import Path
from typing import Any

from aiobale import Client
from aiobale.types import Message

from utils.path_utils import safe_filename


MAX_DOWNLOAD_FILE_SIZE_MB = int(os.getenv("MAX_DOWNLOAD_FILE_SIZE_MB", "10"))
MAX_DOWNLOAD_FILE_BYTES = MAX_DOWNLOAD_FILE_SIZE_MB * 1024 * 1024
MAX_EXTRACTED_CHARS = int(os.getenv("MAX_EXTRACTED_CHARS", "12000"))

TEXT_EXTENSIONS = {
    ".css",
    ".csv",
    ".html",
    ".js",
    ".json",
    ".log",
    ".md",
    ".py",
    ".ts",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
}

ARCHIVE_EXTENSIONS = {".7z", ".rar", ".zip"}
IMAGE_EXTENSIONS = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".svg", ".webp"}
VIDEO_EXTENSIONS = {".avi", ".mkv", ".mov", ".mp4", ".webm"}
AUDIO_EXTENSIONS = {".flac", ".m4a", ".mp3", ".ogg", ".wav"}
UNSUPPORTED_BINARY_EXTENSIONS = (
    ARCHIVE_EXTENSIONS | IMAGE_EXTENSIONS | VIDEO_EXTENSIONS | AUDIO_EXTENSIONS
)


def _document_name(name: Any) -> str | None:
    if isinstance(name, str):
        return name

    if isinstance(name, dict):
        for value in name.values():
            if isinstance(value, str):
                return value

    return None


def _document_mime_type(document: Any, name: str | None) -> str | None:
    mime_type = getattr(document, "mime_type", None)
    if mime_type:
        return mime_type

    guessed, _ = mimetypes.guess_type(name or "")
    return guessed


def _document_caption(document: Any) -> str | None:
    caption = getattr(document, "caption", None)
    if caption is None:
        return None

    return getattr(caption, "content", None)


def _decode_text_file(path: Path) -> tuple[str, str]:
    last_error: UnicodeDecodeError | None = None

    for encoding in ("utf-8", "utf-8-sig", "cp1256", "latin-1"):
        try:
            return path.read_text(encoding=encoding), encoding
        except UnicodeDecodeError as error:
            last_error = error

    if last_error is not None:
        raise last_error

    return path.read_text(encoding="utf-8"), "utf-8"


def _truncate_text(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False

    return text[:max_chars], True


def _extraction_result(
    parser: str,
    extracted_text: str | None = None,
    extraction_error: str | None = None,
    truncated: bool = False,
) -> dict:
    return {
        "extracted_text": extracted_text,
        "extraction_error": extraction_error,
        "truncated": truncated,
        "parser": parser,
    }


def extract_text_from_plain_file(path: Path, max_chars: int) -> dict:
    try:
        text, _encoding = _decode_text_file(path)
        text, truncated = _truncate_text(text, max_chars)
        return _extraction_result(
            parser="plain_text",
            extracted_text=text,
            truncated=truncated,
        )
    except Exception as error:
        return _extraction_result(
            parser="plain_text",
            extraction_error=f"{type(error).__name__}: {error}",
        )


def extract_text_from_pdf(path: Path, max_chars: int) -> dict:
    try:
        from pypdf import PdfReader
    except ImportError:
        return _extraction_result(
            parser="pdf",
            extraction_error="PDF extraction requires optional dependency: pypdf",
        )

    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        total_chars = 0
        truncated = False

        for page in reader.pages:
            page_text = page.extract_text()
            if not page_text:
                continue

            remaining = max_chars - total_chars
            if remaining <= 0:
                truncated = True
                break

            if len(page_text) > remaining:
                parts.append(page_text[:remaining])
                truncated = True
                break

            parts.append(page_text)
            total_chars += len(page_text)

        text = "\n\n".join(part.strip() for part in parts if part.strip())
        text, final_truncated = _truncate_text(text, max_chars)
        truncated = truncated or final_truncated
        if not text:
            return _extraction_result(
                parser="pdf",
                extraction_error=(
                    "No extractable text found in PDF. It may be scanned or image-based."
                ),
            )

        return _extraction_result(
            parser="pdf",
            extracted_text=text,
            truncated=truncated,
        )
    except Exception as error:
        return _extraction_result(
            parser="pdf",
            extraction_error=f"{type(error).__name__}: {error}",
        )


def extract_text_from_docx(path: Path, max_chars: int) -> dict:
    try:
        from docx import Document
    except ImportError:
        return _extraction_result(
            parser="docx",
            extraction_error="DOCX extraction requires optional dependency: python-docx",
        )

    try:
        document = Document(str(path))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    parts.append(" | ".join(cells))

        text, truncated = _truncate_text("\n".join(parts), max_chars)
        if not text:
            return _extraction_result(
                parser="docx",
                extraction_error="No extractable text found in DOCX.",
            )

        return _extraction_result(
            parser="docx",
            extracted_text=text,
            truncated=truncated,
        )
    except Exception as error:
        return _extraction_result(
            parser="docx",
            extraction_error=f"{type(error).__name__}: {error}",
        )


def extract_text_from_xlsx(path: Path, max_chars: int) -> dict:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return _extraction_result(
            parser="xlsx",
            extraction_error="XLSX extraction requires optional dependency: openpyxl",
        )

    try:
        workbook = load_workbook(str(path), data_only=True, read_only=True)
        try:
            parts: list[str] = []
            current_length = 0
            truncated = False

            def append_limited(line: str) -> bool:
                nonlocal current_length, truncated

                separator_length = 1 if parts else 0
                added_length = separator_length + len(line)
                if current_length + added_length <= max_chars:
                    parts.append(line)
                    current_length += added_length
                    return True

                remaining = max_chars - current_length - separator_length
                if remaining > 0:
                    parts.append(line[:remaining])
                    current_length = max_chars
                truncated = True
                return False

            for worksheet in workbook.worksheets:
                if not append_limited(f"Sheet: {worksheet.title}"):
                    break

                for row_index, row in enumerate(
                    worksheet.iter_rows(values_only=True),
                    start=1,
                ):
                    values = [
                        str(value)
                        for value in row
                        if value is not None and str(value) != ""
                    ]
                    if not values:
                        continue

                    if not append_limited(f"row {row_index}: {' | '.join(values)}"):
                        break

                if truncated:
                    break

            text = "\n".join(parts).strip()
        finally:
            workbook.close()

        if not text:
            return _extraction_result(
                parser="xlsx",
                extraction_error="No non-empty cell values found in XLSX.",
            )

        return _extraction_result(
            parser="xlsx",
            extracted_text=text,
            truncated=truncated,
        )
    except Exception as error:
        return _extraction_result(
            parser="xlsx",
            extraction_error=f"{type(error).__name__}: {error}",
        )


def extract_text_from_file(
    path: Path,
    mime_type: str | None = None,
    max_chars: int = MAX_EXTRACTED_CHARS,
) -> dict:
    suffix = path.suffix.lower()

    if suffix in UNSUPPORTED_BINARY_EXTENSIONS:
        return _extraction_result(
            parser="unsupported",
            extraction_error="Unsupported file type for text extraction.",
        )

    if suffix in TEXT_EXTENSIONS or (mime_type or "").startswith("text/"):
        return extract_text_from_plain_file(path, max_chars)

    if suffix == ".pdf" or mime_type == "application/pdf":
        return extract_text_from_pdf(path, max_chars)

    if (
        suffix == ".docx"
        or mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return extract_text_from_docx(path, max_chars)

    if (
        suffix == ".xlsx"
        or mime_type
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    ):
        return extract_text_from_xlsx(path, max_chars)

    return _extraction_result(
        parser="unsupported",
        extraction_error="Unsupported file type for text extraction.",
    )


async def document_to_dict(
    client: Client,
    msg: Message,
    download_dir: Path,
) -> dict | None:
    document = msg.document
    if document is None:
        return None

    name = _document_name(getattr(document, "name", None))
    safe_name = safe_filename(name or f"file_{document.file_id}")
    local_path = download_dir / f"{msg.message_id}_{safe_name}"
    size = getattr(document, "size", None)
    mime_type = _document_mime_type(document, name)

    result = {
        "file_id": document.file_id,
        "access_hash": document.access_hash,
        "name": name,
        "size": size,
        "mime_type": mime_type,
        "caption": _document_caption(document),
        "local_path": str(local_path),
        "downloaded": False,
        "extracted_text": None,
        "extraction_error": None,
        "truncated": False,
        "parser": None,
    }

    if size is not None and size > MAX_DOWNLOAD_FILE_BYTES:
        result["extraction_error"] = (
            f"File is larger than MAX_DOWNLOAD_FILE_SIZE_MB "
            f"({MAX_DOWNLOAD_FILE_SIZE_MB} MB); download skipped."
        )
        return result

    try:
        download_dir.mkdir(parents=True, exist_ok=True)
        await client.download_file(
            file_id=document.file_id,
            access_hash=document.access_hash,
            destination=local_path,
        )
        result["downloaded"] = True
    except Exception as error:
        result["extraction_error"] = f"Download failed: {type(error).__name__}: {error}"
        return result

    extraction = extract_text_from_file(
        path=local_path,
        mime_type=mime_type,
        max_chars=MAX_EXTRACTED_CHARS,
    )
    result.update(extraction)
    return result
