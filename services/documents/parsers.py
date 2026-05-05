from pathlib import Path

from services.documents.constants import (
    MAX_EXTRACTED_CHARS,
    TEXT_EXTENSIONS,
    UNSUPPORTED_BINARY_EXTENSIONS,
)


def decode_text_file(path: Path) -> tuple[str, str]:
    last_error: UnicodeDecodeError | None = None

    for encoding in ("utf-8", "utf-8-sig", "cp1256", "latin-1"):
        try:
            return path.read_text(encoding=encoding), encoding
        except UnicodeDecodeError as error:
            last_error = error

    if last_error is not None:
        raise last_error

    return path.read_text(encoding="utf-8"), "utf-8"


def truncate_text(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False

    return text[:max_chars], True


def extraction_result(
    parser: str,
    extracted_text: str | None = None,
    extraction_error: str | None = None,
    truncated: bool = False,
) -> dict:
    return {
        "extracted_text": extracted_text,
        "extraction_error": extraction_error,
        "truncated": truncated,
        "parser": parser,
    }


def unsupported_extraction_result() -> dict:
    return extraction_result(
        parser="unsupported",
        extraction_error="Unsupported file type for text extraction.",
    )


def extract_text_from_plain_file(path: Path, max_chars: int) -> dict:
    try:
        text, _encoding = decode_text_file(path)
        text, truncated = truncate_text(text, max_chars)
        return extraction_result(
            parser="plain_text",
            extracted_text=text,
            truncated=truncated,
        )
    except Exception as error:
        return extraction_result(
            parser="plain_text",
            extraction_error=f"{type(error).__name__}: {error}",
        )


def extract_text_from_pdf(path: Path, max_chars: int) -> dict:
    try:
        from pypdf import PdfReader
    except ImportError:
        return extraction_result(
            parser="pdf",
            extraction_error="PDF extraction requires optional dependency: pypdf",
        )

    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        total_chars = 0
        truncated = False

        for page in reader.pages:
            page_text = page.extract_text()
            if not page_text:
                continue

            remaining = max_chars - total_chars
            if remaining <= 0:
                truncated = True
                break

            if len(page_text) > remaining:
                parts.append(page_text[:remaining])
                truncated = True
                break

            parts.append(page_text)
            total_chars += len(page_text)

        text = "\n\n".join(part.strip() for part in parts if part.strip())
        text, final_truncated = truncate_text(text, max_chars)
        truncated = truncated or final_truncated
        if not text:
            return extraction_result(
                parser="pdf",
                extraction_error=(
                    "No extractable text found in PDF. It may be scanned or image-based."
                ),
            )

        return extraction_result(
            parser="pdf",
            extracted_text=text,
            truncated=truncated,
        )
    except Exception as error:
        return extraction_result(
            parser="pdf",
            extraction_error=f"{type(error).__name__}: {error}",
        )


def extract_text_from_docx(path: Path, max_chars: int) -> dict:
    try:
        from docx import Document
    except ImportError:
        return extraction_result(
            parser="docx",
            extraction_error="DOCX extraction requires optional dependency: python-docx",
        )

    try:
        document = Document(str(path))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    parts.append(" | ".join(cells))

        text, truncated = truncate_text("\n".join(parts), max_chars)
        if not text:
            return extraction_result(
                parser="docx",
                extraction_error="No extractable text found in DOCX.",
            )

        return extraction_result(
            parser="docx",
            extracted_text=text,
            truncated=truncated,
        )
    except Exception as error:
        return extraction_result(
            parser="docx",
            extraction_error=f"{type(error).__name__}: {error}",
        )


def extract_text_from_xlsx(path: Path, max_chars: int) -> dict:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return extraction_result(
            parser="xlsx",
            extraction_error="XLSX extraction requires optional dependency: openpyxl",
        )

    try:
        workbook = load_workbook(str(path), data_only=True, read_only=True)
        try:
            parts: list[str] = []
            current_length = 0
            truncated = False

            def append_limited(line: str) -> bool:
                nonlocal current_length, truncated

                separator_length = 1 if parts else 0
                added_length = separator_length + len(line)
                if current_length + added_length <= max_chars:
                    parts.append(line)
                    current_length += added_length
                    return True

                remaining = max_chars - current_length - separator_length
                if remaining > 0:
                    parts.append(line[:remaining])
                    current_length = max_chars
                truncated = True
                return False

            for worksheet in workbook.worksheets:
                if not append_limited(f"Sheet: {worksheet.title}"):
                    break

                for row_index, row in enumerate(
                    worksheet.iter_rows(values_only=True),
                    start=1,
                ):
                    values = [
                        str(value)
                        for value in row
                        if value is not None and str(value) != ""
                    ]
                    if not values:
                        continue

                    if not append_limited(f"row {row_index}: {' | '.join(values)}"):
                        break

                if truncated:
                    break

            text = "\n".join(parts).strip()
        finally:
            workbook.close()

        if not text:
            return extraction_result(
                parser="xlsx",
                extraction_error="No non-empty cell values found in XLSX.",
            )

        return extraction_result(
            parser="xlsx",
            extracted_text=text,
            truncated=truncated,
        )
    except Exception as error:
        return extraction_result(
            parser="xlsx",
            extraction_error=f"{type(error).__name__}: {error}",
        )


def extract_text_from_file(
    path: Path,
    mime_type: str | None = None,
    max_chars: int = MAX_EXTRACTED_CHARS,
) -> dict:
    suffix = path.suffix.lower()

    if suffix in UNSUPPORTED_BINARY_EXTENSIONS:
        return unsupported_extraction_result()

    if suffix in TEXT_EXTENSIONS or (mime_type or "").startswith("text/"):
        return extract_text_from_plain_file(path, max_chars)

    if suffix == ".pdf" or mime_type == "application/pdf":
        return extract_text_from_pdf(path, max_chars)

    if (
        suffix == ".docx"
        or mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return extract_text_from_docx(path, max_chars)

    if (
        suffix == ".xlsx"
        or mime_type
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    ):
        return extract_text_from_xlsx(path, max_chars)

    return unsupported_extraction_result()
